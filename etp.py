import streamlit as st 
from docx import Document

documento = Document("ANOTAR.docx")

for paragrafo in documento.paragraphs:
    if "marcos" in paragrafo.text:
        paragrafo.text=paragrafo.text.replace("marcos", "cerenilo")

    st.write(paragrafo.text)
    
documento.add_paragraph("Mais um paragrafo aqui")


documento.save("Novo_arquivo.docx")